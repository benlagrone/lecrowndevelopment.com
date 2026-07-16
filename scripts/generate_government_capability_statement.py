#!/usr/bin/env python3
"""Generate LeCrown's one-page government capability statement and editable DOCX."""

from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
PDF = ROOT / "output" / "pdf" / "lecrown-development-capability-statement.pdf"
PUBLIC_PDF = ROOT / "public" / "downloads" / PDF.name
DOCX = ROOT / "output" / "docx" / "lecrown-development-capability-statement.docx"
LOGO = ROOT / "public" / "logo-mark.png"

NAVY = colors.HexColor("#0F1726")
BLUE = colors.HexColor("#073C57")
TEAL = colors.HexColor("#0B8790")
LIME = colors.HexColor("#B8D328")
PALE = colors.HexColor("#EAF5F3")
CREAM = colors.HexColor("#F7F4EC")
MUTED = colors.HexColor("#4F5D6B")
LINE = colors.HexColor("#CCD7D8")

CORE = [
    "Enterprise data and AI platform architecture, governance, MLOps and knowledge systems",
    "Cloud modernization, platform engineering, DevSecOps, infrastructure as code and FinOps",
    "Custom software, portals, APIs, workflow automation, systems integration and decision support",
    "Technology strategy, solution architecture, program delivery and operating-model design",
]

EXPERIENCE = [
    ("Sysco | Enterprise Data & AI Platform Architecture (Key Personnel, 2024-present)",
     "Designed governed cloud-platform patterns, self-service provisioning, identity and policy controls, reusable delivery capabilities and operating-model guidance for enterprise data and AI teams."),
    ("LeCrown Development | Secure Digital Portals & Automation (2025-present)",
     "Architected and delivered client portals, authenticated document workflows, intake automation, service integrations and containerized cloud deployments from discovery through production."),
    ("Energy & Industrial Enterprises | Cloud, Data & Application Delivery (Key Personnel)",
     "Led solution architecture, software delivery, integration and modernization initiatives in asset-intensive, regulated operating environments with globally distributed stakeholders."),
]

DIFFERENTIATORS = [
    "50% cycle and cost reduction achieved through automated data-pipeline delivery.",
    "$20M+ in program value protected through architecture and quality governance.",
    "30% delivery improvement achieved through Agile and DevOps transformation.",
    "75% cost savings achieved through targeted business-process automation.",
]

DATA = [
    ("NAICS", "541512 (Primary)  |  541511  |  541519  |  541611"),
    ("PSC", "DA01  |  DD01"),
    ("NIGP", "918-71  |  920-40  |  920-45  |  920-37"),
    ("Service Area", "Nationwide delivery | Houston-based"),
    ("Delivery", "Remote, hybrid and on-site | Prime or subcontractor"),
]


def p(text, style):
    return Paragraph(escape(text), style)


def styles():
    return {
        "title": ParagraphStyle("title", fontName="Helvetica-Bold", fontSize=16, leading=17, textColor=NAVY),
        "company": ParagraphStyle("company", fontName="Times-Bold", fontSize=18, leading=19, textColor=NAVY),
        "tag": ParagraphStyle("tag", fontName="Helvetica", fontSize=7.5, leading=9.5, textColor=MUTED),
        "contact": ParagraphStyle("contact", fontName="Helvetica", fontSize=7.2, leading=10, textColor=MUTED, alignment=TA_RIGHT),
        "section": ParagraphStyle("section", fontName="Helvetica-Bold", fontSize=9.2, leading=11, textColor=colors.white),
        "body": ParagraphStyle("body", fontName="Helvetica", fontSize=8.1, leading=11.2, textColor=NAVY),
        "bullet": ParagraphStyle("bullet", fontName="Helvetica", fontSize=7.8, leading=10.7, leftIndent=9, firstLineIndent=-7, textColor=NAVY),
        "exp_title": ParagraphStyle("exp_title", fontName="Helvetica-Bold", fontSize=7.5, leading=9.4, textColor=BLUE),
        "exp_body": ParagraphStyle("exp_body", fontName="Helvetica", fontSize=7.2, leading=9.5, textColor=MUTED),
        "data_label": ParagraphStyle("data_label", fontName="Helvetica-Bold", fontSize=7.2, leading=8.8, textColor=BLUE),
        "data_value": ParagraphStyle("data_value", fontName="Helvetica", fontSize=7.0, leading=8.8, textColor=NAVY),
        "footer": ParagraphStyle("footer", fontName="Helvetica", fontSize=6.5, leading=8.2, textColor=MUTED),
    }


def section_header(label, width, s):
    t = Table([[p(label.upper(), s["section"])]], colWidths=[width])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), BLUE),
        ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


def bullet_list(items, s):
    return [p(f"•  {item}", s["bullet"]) for item in items]


def build_pdf():
    s = styles()
    doc = SimpleDocTemplate(str(PDF), pagesize=letter, leftMargin=0.36*inch, rightMargin=0.36*inch,
                            topMargin=0.3*inch, bottomMargin=0.27*inch,
                            title="LeCrown Development Capability Statement", author="LeCrown Development")
    content_width = 7.78 * inch
    logo = Image(str(LOGO), width=0.62*inch, height=0.62*inch)
    identity = [p("LeCrown Development", s["company"]), p("ENTERPRISE AI • DATA PLATFORMS • CLOUD ENGINEERING • DIGITAL SYSTEMS", s["tag"])]
    contact = Paragraph("Benjamin LaGrone  |  Founder &amp; Principal Technology Consultant<br/>"
                        "Houston, TX 77024  |  (910) 236-9853<br/>"
                        "benjamin.lagrone@lecrowndevelopment.com  |  lecrowndevelopment.com", s["contact"])
    head = Table([[logo, identity, contact]], colWidths=[0.72*inch, 3.55*inch, 3.51*inch])
    head.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 0),
                              ("RIGHTPADDING", (0, 0), (-1, -1), 0), ("TOPPADDING", (0, 0), (-1, -1), 0),
                              ("BOTTOMPADDING", (0, 0), (-1, -1), 0)]))
    story = [head, Spacer(1, 0.08*inch)]
    title = Table([[p("CAPABILITY STATEMENT", s["title"]), p("GOVERNMENT & PRIME CONTRACTOR DELIVERY", s["contact"])]],
                  colWidths=[4.4*inch, 3.38*inch])
    title.setStyle(TableStyle([("LINEABOVE", (0, 0), (-1, -1), 2, LIME), ("LINEBELOW", (0, 0), (-1, -1), 0.6, LINE),
                               ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 0),
                               ("RIGHTPADDING", (0, 0), (-1, -1), 0), ("TOPPADDING", (0, 0), (-1, -1), 6),
                               ("BOTTOMPADDING", (0, 0), (-1, -1), 6)]))
    story += [title, Spacer(1, 0.1*inch)]

    intro = p("LeCrown Development designs and delivers governed enterprise technology that moves complex initiatives from strategy to secure, measurable operation. Founder-led execution combines architecture, engineering and delivery leadership in one accountable team.", s["body"])

    left_w, right_w = 4.83*inch, 2.83*inch
    core = [section_header("Core Competencies", left_w, s), Spacer(1, 0.05*inch), intro, Spacer(1, 0.05*inch)] + bullet_list(CORE, s)
    data_rows = [[p(k, s["data_label"]), p(v, s["data_value"])] for k, v in DATA]
    data_table = Table(data_rows, colWidths=[0.66*inch, 2.05*inch])
    data_table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LINEBELOW", (0, 0), (-1, -2), 0.35, LINE),
                                    ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                                    ("TOPPADDING", (0, 0), (-1, -1), 3.5), ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5)]))
    company_data = [section_header("Company Data", right_w, s), Spacer(1, 0.05*inch), data_table,
                    Spacer(1, 0.06*inch), p("Credentials", s["exp_title"]),
                    p("Azure Data Fundamentals • Azure Databricks Platform Architect • Scrum Master Certified • B.S., University of Houston", s["exp_body"])]
    top = Table([[core, company_data]], colWidths=[left_w, right_w], rowHeights=[2.72*inch], hAlign="LEFT")
    top.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (0, 0), 0),
                             ("RIGHTPADDING", (0, 0), (0, 0), 7), ("LEFTPADDING", (1, 0), (1, 0), 7),
                             ("RIGHTPADDING", (1, 0), (1, 0), 0), ("LINEBEFORE", (1, 0), (1, 0), 0.5, LINE),
                             ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 0)]))
    story += [top, Spacer(1, 0.11*inch), section_header("Company Experience", content_width, s), Spacer(1, 0.05*inch)]
    exp_cells = []
    for title_text, body_text in EXPERIENCE:
        exp_cells.append([p(title_text, s["exp_title"]), p(body_text, s["exp_body"])])
    exp = Table([exp_cells], colWidths=[content_width/3]*3, rowHeights=[1.32*inch])
    exp.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("BACKGROUND", (0, 0), (-1, -1), colors.white),
                             ("BOX", (0, 0), (-1, -1), 0.45, LINE), ("INNERGRID", (0, 0), (-1, -1), 0.45, LINE),
                             ("LEFTPADDING", (0, 0), (-1, -1), 7), ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                             ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6)]))
    story += [exp, Spacer(1, 0.11*inch), section_header("Differentiators — Documented Delivery Outcomes", content_width, s), Spacer(1, 0.05*inch)]
    diff_cells = []
    for item in DIFFERENTIATORS:
        value, text = item.split(" ", 1)
        diff_cells.append([p(value, ParagraphStyle("metric", parent=s["title"], fontSize=15, leading=16, textColor=TEAL)), p(text, s["exp_body"])])
    diff = Table([diff_cells], colWidths=[content_width/4]*4, rowHeights=[0.82*inch])
    diff.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("BACKGROUND", (0, 0), (-1, -1), PALE),
                              ("BOX", (0, 0), (-1, -1), 0.5, TEAL), ("INNERGRID", (0, 0), (-1, -1), 0.4, LINE),
                              ("LEFTPADDING", (0, 0), (-1, -1), 7), ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                              ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6)]))
    readiness = Table([
        [p("TECHNOLOGY ENVIRONMENT", s["data_label"]), p("DELIVERY READINESS", s["data_label"])],
        [p("GCP • Azure • AWS • Databricks • Snowflake • Kubernetes • Terraform • ServiceNow • Kafka • Python • FastAPI • React", s["exp_body"]),
         p("Discovery • Architecture • MVP/Pilot • Modernization • Integration • Documentation • Production deployment • Knowledge transfer", s["exp_body"])],
    ], colWidths=[content_width/2]*2, rowHeights=[0.24*inch, 0.52*inch])
    readiness.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("BACKGROUND", (0, 0), (-1, -1), CREAM),
                                   ("BOX", (0, 0), (-1, -1), 0.45, LINE), ("LINEBEFORE", (1, 0), (1, -1), 0.45, LINE),
                                   ("LEFTPADDING", (0, 0), (-1, -1), 7), ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                                   ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
    story += [diff, Spacer(1, 0.1*inch), readiness, Spacer(1, 0.09*inch), p("Selected outcomes reflect documented key-personnel engagements. References and supporting detail are available for qualified opportunities. LeCrown Development is available for prime, subcontract and teaming engagements.", s["footer"])]
    PDF.parent.mkdir(parents=True, exist_ok=True)
    PUBLIC_PDF.parent.mkdir(parents=True, exist_ok=True)
    def draw_page(canvas, _doc):
        canvas.saveState()
        canvas.setFillColor(colors.white)
        canvas.rect(0, 0, letter[0], letter[1], fill=1, stroke=0)
        canvas.setFillColor(BLUE)
        canvas.rect(0, 0, letter[0], 0.2*inch, fill=1, stroke=0)
        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica-Bold", 6.3)
        canvas.drawCentredString(letter[0] / 2, 0.075*inch, "LECROWN DEVELOPMENT  •  FOUNDER-LED ENTERPRISE TECHNOLOGY DELIVERY  •  LECROWNDEVELOPMENT.COM")
        canvas.restoreState()

    doc.build(story, onFirstPage=draw_page)
    PUBLIC_PDF.write_bytes(PDF.read_bytes())


def set_cell(cell, text, bold=False, color="0F1726", size=8):
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell(cell, text.upper(), True, "FFFFFF", 9)
    cell._tc.get_or_add_tcPr().append(_shade("073C57"))


def _shade(fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    return shd


def build_docx():
    d = Document()
    sec = d.sections[0]
    sec.page_height, sec.page_width = Inches(11), Inches(8.5)
    sec.top_margin = sec.bottom_margin = Inches(0.32)
    sec.left_margin = sec.right_margin = Inches(0.38)
    normal = d.styles["Normal"]
    normal.font.name, normal.font.size = "Arial", Pt(8)
    header = d.add_table(rows=1, cols=3)
    header.autofit = False
    header.columns[0].width, header.columns[1].width, header.columns[2].width = Inches(.7), Inches(3.25), Inches(3.55)
    header.cell(0, 0).paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(.45))
    brand_cell = header.cell(0, 1)
    brand_cell.text = ""
    brand_name = brand_cell.paragraphs[0].add_run("LeCrown Development")
    brand_name.bold, brand_name.font.name, brand_name.font.size = True, "Arial", Pt(14)
    brand_name.font.color.rgb = RGBColor.from_string("0F1726")
    brand_tag = brand_cell.add_paragraph().add_run("ENTERPRISE AI • DATA PLATFORMS • CLOUD ENGINEERING")
    brand_tag.bold, brand_tag.font.name, brand_tag.font.size = True, "Arial", Pt(7.5)
    brand_tag.font.color.rgb = RGBColor.from_string("4F5D6B")
    set_cell(header.cell(0, 2), "Benjamin LaGrone | Founder & Principal Technology Consultant\nHouston, TX 77024 | (910) 236-9853\nbenjamin.lagrone@lecrowndevelopment.com | lecrowndevelopment.com", False, "4F5D6B", 7)
    header.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title = d.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("CAPABILITY STATEMENT")
    r.bold, r.font.name, r.font.size, r.font.color.rgb = True, "Arial", Pt(17), RGBColor.from_string("0F1726")
    add_heading(d, "Core Competencies")
    d.add_paragraph("LeCrown Development designs and delivers governed enterprise technology that moves complex initiatives from strategy to secure, measurable operation. Founder-led execution combines architecture, engineering and delivery leadership in one accountable team.")
    for item in CORE:
        para = d.add_paragraph(style=None)
        para.paragraph_format.left_indent = Inches(.12)
        para.paragraph_format.first_line_indent = Inches(-.1)
        para.add_run("• " + item)
    add_heading(d, "Company Experience")
    exp = d.add_table(rows=1, cols=3)
    for i, (heading, body) in enumerate(EXPERIENCE):
        set_cell(exp.cell(0, i), heading + "\n" + body, False, "0F1726", 7)
    add_heading(d, "Differentiators — Documented Delivery Outcomes")
    dif = d.add_table(rows=1, cols=4)
    for i, item in enumerate(DIFFERENTIATORS):
        set_cell(dif.cell(0, i), item, True, "073C57", 8)
        dif.cell(0, i)._tc.get_or_add_tcPr().append(_shade("EAF5F3"))
    add_heading(d, "Company Data")
    data = d.add_table(rows=len(DATA), cols=2)
    for i, (k, v) in enumerate(DATA):
        set_cell(data.cell(i, 0), k, True, "073C57", 7)
        set_cell(data.cell(i, 1), v, False, "0F1726", 7)
    d.add_paragraph("Credentials: Azure Data Fundamentals • Azure Databricks Platform Architect • Scrum Master Certified • B.S., University of Houston")
    foot = d.add_paragraph("Selected outcomes reflect documented key-personnel engagements. References and supporting detail are available for qualified opportunities. LeCrown Development is available for prime, subcontract and teaming engagements.")
    foot.runs[0].font.size = Pt(6)
    foot.runs[0].font.color.rgb = RGBColor.from_string("4F5D6B")
    DOCX.parent.mkdir(parents=True, exist_ok=True)
    d.save(DOCX)


if __name__ == "__main__":
    build_pdf()
    build_docx()
    print(PDF)
    print(DOCX)
    print(PUBLIC_PDF)
